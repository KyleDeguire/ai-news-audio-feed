#!/usr/bin/env python3
# scripts/compose_transcript.py
# Build HTML + DOCX for the transcript with robust fallbacks and sane paragraphing.

import json, re, html
from pathlib import Path
from typing import Any, Dict, List
from datetime import datetime
import pytz
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

AUDIO_DIR = Path("audio")
HEADINGS = [
    "New Products & Capabilities",
    "Strategic Business Impact",
    "Implementation Opportunities",
    "Market Dynamics",
    "Talent Market Shifts",
]

# ---------- time helpers ----------
def denver_now() -> datetime:
    return datetime.now(pytz.timezone("America/Denver"))

def subject_date(dt: datetime) -> str:
    return dt.strftime("%d %b %y")  # e.g., 10 Nov 25

# ---------- locate latest assets ----------
def latest(prefix: str, ext: str) -> Path | None:
    files = sorted(AUDIO_DIR.glob(f"{prefix}_*.{ext}"), key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0] if files else None

def latest_json() -> Path | None: return latest("ai_news","json")
def latest_txt()  -> Path | None: return latest("ai_news","txt")

# ---------- load ----------
def load_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8", errors="ignore"))

def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")

# ---------- normalize model JSON ----------
def normalize_from_json(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Target shape:
    {
      "sections":[
        {"title":"...", "paragraphs":[{"text":"...", "sources":[1,2]}, ...]}
      ],
      "footnotes":[{"id":1,"title":"...","url":"..."}],
      "spoken":"..."   # optional
    }
    """
    out = {"sections": [], "footnotes": [], "spoken": (data.get("spoken") or "").strip()}

    # footnotes
    fns = data.get("footnotes") or []
    cleaned = []
    for i, f in enumerate(fns, start=1):
        if not isinstance(f, dict): continue
        fid = f.get("id") if isinstance(f.get("id"), int) else i
        ttl = (f.get("title") or "").strip()
        url = (f.get("url") or "").strip()
        if url:
            cleaned.append({"id": fid, "title": ttl, "url": url})
    out["footnotes"] = cleaned

    # sections
    secs = data.get("sections") or []
    norm_secs = []
    for s in secs:
        if isinstance(s, str):
            norm_secs.append({"title": s.strip(), "paragraphs": []})
            continue
        if not isinstance(s, dict): 
            continue
        title = (s.get("title") or "").strip()
        paras = []
        for p in (s.get("paragraphs") or []):
            if isinstance(p, str):
                t = p.strip()
                if t: paras.append({"text": t, "sources": []})
                continue
            if isinstance(p, dict):
                t = (p.get("text") or "").strip()
                srcs = [int(x) for x in p.get("sources") or [] if str(x).isdigit()]
                if t: paras.append({"text": t, "sources": srcs})
        norm_secs.append({"title": title, "paragraphs": paras})
    out["sections"] = norm_secs
    return out

# ---------- paragraph builders ----------
_SENT_SPLIT = re.compile(r'(?<=[.!?])\s+(?=[A-Z(])')

def paragraphs_from_free_text(text: str) -> List[str]:
    """Make readable paragraphs from a long script:
       - ignore trailing 'Sources:' block if present
       - group 3–5 sentences per paragraph
    """
    if not text: return []
    # strip any "Sources:" tail
    lower = text.lower()
    cut = lower.find("\nsources:")
    if cut != -1:
        text = text[:cut]

    # collapse internal whitespace
    text = re.sub(r'[ \t]+', ' ', text.strip())
    # if author already used blank lines, keep them
    blocks = [b.strip() for b in re.split(r'\n\s*\n', text) if b.strip()]
    if len(blocks) >= 3:
        return blocks

    # otherwise split to sentences and batch into 3–5 per paragraph
    sents = [s.strip() for s in _SENT_SPLIT.split(text) if s.strip()]
    out, buf = [], []
    for s in sents:
        buf.append(s)
        if len(buf) >= 4:  # target 4 sentences per paragraph
            out.append(" ".join(buf)); buf = []
    if buf: out.append(" ".join(buf))
    return out

def distribute_across_headings(paras: List[str]) -> List[Dict[str, Any]]:
    """Map paragraphs into the five standard headings evenly."""
    if not paras: return [{"title": t, "paragraphs": []} for t in HEADINGS]
    buckets = [[] for _ in HEADINGS]
    for i, p in enumerate(paras):
        buckets[i % len(HEADINGS)].append({"text": p, "sources": []})
    return [{"title": HEADINGS[i], "paragraphs": buckets[i]} for i in range(len(HEADINGS))]

# ---------- render HTML ----------
def build_html(sections: List[Dict[str, Any]], footnotes: List[Dict[str, Any]]) -> str:
    parts: List[str] = []
    parts.append("<div style=\"font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,Arial,sans-serif;font-size:15px;line-height:1.6;\">")
    for sec in sections:
        title = (sec.get("title") or "").strip()
        if title:
            parts.append(f"<p style='margin:0 0 8px 0'><strong>{html.escape(title.upper())}</strong></p>")
        for para in sec.get("paragraphs", []):
            text = (para.get("text") or "").strip()
            srcs = para.get("sources") or []
            sup = f"<sup>{','.join(str(int(s)) for s in srcs if str(s).isdigit())}</sup>" if srcs else ""
            if text:
                parts.append(f"<p style='margin:0 0 18px 0'>{html.escape(text)}{sup}</p>")
    if footnotes:
        parts.append("<hr style='margin:8px 0;border:none;border-top:1px solid #e4e7ef'>")
        parts.append("<p style='margin:0 0 6px 0'><strong>Sources</strong></p><ul style='margin:0 0 0 18px;padding:0'>")
        for f in footnotes:
            iid = f.get("id"); ttl = (f.get("title") or "").strip(); url = (f.get("url") or "").strip()
            if not url: continue
            label = f"[{iid}]" if iid is not None else "-"
            ttl_esc = html.escape(ttl) if ttl else ""
            url_esc = html.escape(url)
            if ttl_esc:
                parts.append(f"<li>{label} {ttl_esc} — <a href=\"{url_esc}\">{url_esc}</a></li>")
            else:
                parts.append(f"<li>{label} <a href=\"{url_esc}\">{url_esc}</a></li>")
        parts.append("</ul>")
    parts.append("</div>")
    return "".join(parts)

# ---------- render DOCX ----------
def build_docx(path: Path, sections: List[Dict[str, Any]], footnotes: List[Dict[str, Any]]):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.font.size = Pt(11)

    def space(p):
        pf = p.paragraph_format
        pf.space_after = Pt(18)
        pf.line_spacing = 1.2

    for sec in sections:
        title = (sec.get("title") or "").strip()
        if title:
            p = doc.add_paragraph(); r = p.add_run(title); r.bold = True; space(p)
        for para in sec.get("paragraphs", []):
            t = (para.get("text") or "").strip()
            if not t: continue
            srcs = para.get("sources") or []
            p = doc.add_paragraph(); r = p.add_run(t)
            if srcs:
                r2 = p.add_run(" " + ",".join(str(int(s)) for s in srcs if str(s).isdigit()))
                r2.font.superscript = True
            space(p)

    if footnotes:
        p = doc.add_paragraph(); r = p.add_run("Sources"); r.bold = True; space(p)
        for f in footnotes:
            iid = f.get("id"); ttl = (f.get("title") or "").strip(); url = (f.get("url") or "").strip()
            if not url: continue
            line = f"[{iid}] {ttl} — {url}" if ttl else f"[{iid}] {url}"
            p = doc.add_paragraph(line); space(p)

    doc.save(path)

# ---------- main ----------
def main():
    jpath = latest_json()
    if not jpath:
        raise SystemExit("No ai_news_*.json found. Run generate_brief.py first.")
    data = load_json(jpath)
    norm = normalize_from_json(data)

    # Preferred: real paragraphs from JSON sections
    has_paras = any(len(sec.get("paragraphs") or []) > 0 for sec in norm["sections"])
    if not has_paras:
        # Fallback 1: model provided a single 'spoken' string
        spoken = norm.get("spoken") or ""
        if not spoken:
            # Fallback 2: transcript .txt
            tpath = latest_txt()
            if tpath:
                spoken = load_txt(tpath)
        paras = paragraphs_from_free_text(spoken)
        norm["sections"] = distribute_across_headings(paras)

    sections = norm["sections"]
    footnotes = norm["footnotes"]

    today = denver_now()
    base = f"AI Exec Brief (transcript) - {subject_date(today)}"
    html_path = AUDIO_DIR / f"{base}.html"
    docx_path = AUDIO_DIR / f"{base}.docx"
    subject = base

    html_body = build_html(sections, footnotes)
    html_path.write_text(html_body, encoding="utf-8")
    build_docx(docx_path, sections, footnotes)

    print(f"HTML_BODY={html_path}")
    print(f"DOCX_PATH={docx_path}")
    print(f"SUBJECT_LINE={subject}")

if __name__ == "__main__":
    main()
