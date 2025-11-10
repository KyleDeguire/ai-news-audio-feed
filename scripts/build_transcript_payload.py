#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os, re, base64, datetime, pytz, sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
AUDIO_DIR = ROOT / "audio"
OUT_DIR   = ROOT / "transcripts"
OUT_DIR.mkdir(parents=True, exist_ok=True)

def denver_stamp_today():
    tz = pytz.timezone("America/Denver")
    d = datetime.datetime.now(tz).date()
    return d.strftime("%Y%m%d")

def friendly_from_stamp(stamp):
    d = datetime.datetime.strptime(stamp, "%Y%m%d")
    return d.strftime("%d %b %y")  # e.g., 27 Sep 25

def find_latest_transcript():
    today = denver_stamp_today()
    preferred = AUDIO_DIR / f"ai_news_{today}.txt"
    if preferred.exists():
        return preferred, today
    cands = sorted(AUDIO_DIR.glob("ai_news_*.txt"))
    if not cands:
        raise FileNotFoundError("No transcript files found under /audio/")
    latest = cands[-1]
    m = re.search(r"(\d{8})", latest.stem)
    if not m:
        raise RuntimeError(f"Cannot parse stamp from {latest.name}")
    return latest, m.group(1)

def load_and_split(txt_path: Path):
    text = txt_path.read_text(encoding="utf-8").strip()
    # Your generator writes:
    # <spoken text>\n---\nSources:\n[n] Title --- URL
    parts = re.split(r"^\s*---\s*\n\s*Sources\s*:\s*$", text, flags=re.IGNORECASE|re.MULTILINE)
    if len(parts) != 2:
        # also accept single-line divider
        parts = re.split(r"^---\s*Sources\s*:\s*$", text, flags=re.IGNORECASE|re.MULTILINE)
    if len(parts) != 2:
        raise ValueError("Transcript must contain a '---' divider and a following 'Sources:' block.")
    body, sources_raw = parts[0].strip(), parts[1].strip()

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body) if p.strip()]

    sources = []
    for line in sources_raw.splitlines():
        line = line.strip()
        if not line:
            continue
        m = re.match(r"\[(\d+)\]\s*(.+?)\s*---\s*(\S+)", line)
        if m:
            num, title, url = int(m.group(1)), m.group(2).strip(), m.group(3).strip()
            sources.append((num, title, url))
    sources.sort(key=lambda x: x[0])
    return paragraphs, sources

def sentence_chunks_with_cites(paragraph, sources):
    """
    Returns [(sentence_text, [citations])].
    If [n] markers exist in text, use them. Otherwise, heuristic keyword match from source titles.
    """
    sentences = re.split(r'(?<=[\.\?\!])\s+', paragraph)
    out = []

    # Build simple keyword sets from source titles
    kw = []
    for num, title, _ in sources:
        tokens = [t for t in re.split(r"[\s:/\-]+", title) if len(t) >= 4]
        kw.append((num, set(t.lower() for t in tokens)))

    for s in sentences:
        s = s.strip()
        if not s:
            continue
        direct = [int(n) for n in re.findall(r"\[(\d+)\]", s)]
        s_text = re.sub(r"\s*\[(\d+)\]\s*", "", s).strip()
        cites = sorted(set(direct))
        if not cites and kw:
            words = set(re.findall(r"[A-Za-z0-9]+", s_text.lower()))
            for num, toks in kw:
                if words & toks:
                    cites.append(num)
        out.append((s_text, sorted(set(cites))))
    return out

def build_docx(paragraphs, sources, out_docx: Path):
    doc = Document()
    # Body paragraphs with spacing and superscripts
    for para in paragraphs:
        chunks = sentence_chunks_with_cites(para, sources)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(18)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.2

        first = True
        for text, cites in chunks:
            if not first:
                p.add_run(" ")
            first = False
            p.add_run(text)
            for c in cites:
                sup = p.add_run(f" {c}")
                sup.font.superscript = True

    # Sources heading
    doc.add_paragraph()
    hdr = doc.add_paragraph()
    run = hdr.add_run("Sources:")
    run.bold = True

    # Sources list (compact)
    for num, title, url in sources:
        li = doc.add_paragraph(f"[{num}] {title} — {url}")
        li.paragraph_format.space_after = Pt(6)

    doc.save(out_docx)

def html_escape(s: str) -> str:
    return (s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"))

def build_html(paragraphs, sources):
    html_paras = []
    for para in paragraphs:
        chunks = sentence_chunks_with_cites(para, sources)
        segs = []
        for text, cites in chunks:
            frag = html_escape(text)
            if cites:
                frag += "".join(f"<sup>{c}</sup>" for c in cites)
            segs.append(frag)
        html_paras.append(f"<p style='margin:0 0 18pt; line-height:1.2;'>{' '.join(segs)}</p>")

    items = "".join(
        f"<li>[{num}] {html_escape(title)} — <a href='{url}'>{url}</a></li>"
        for num, title, url in sources
    )
    html_sources = (
        "<p style='margin:18pt 0 6pt; line-height:1.2;'><strong>Sources:</strong></p>"
        f"<ul style='margin:0 0 12pt 18pt; line-height:1.2;'>{items}</ul>"
    )
    return "\n".join(html_paras) + "\n" + html_sources

def main():
    txt_path, stamp = find_latest_transcript()
    paragraphs, sources = load_and_split(txt_path)

    friendly = friendly_from_stamp(stamp)  # e.g., 27 Sep 25
    subject  = f"AI Exec Brief (transcript) - {friendly}"
    docx_path = OUT_DIR / f"{subject}.docx"

    build_docx(paragraphs, sources, docx_path)
    html = build_html(paragraphs, sources)
    b64 = base64.b64encode(html.encode("utf-8")).decode("ascii")

    # Surface outputs to GitHub Actions
    gh_out = os.environ.get("GITHUB_OUTPUT")
    lines = [
        f"SUBJECT={subject}",
        f"ATTACH_PATH={docx_path.as_posix()}",
        f"BODY_B64={b64}",
    ]
    if gh_out:
        with open(gh_out, "a", encoding="utf-8") as f:
            for line in lines:
                f.write(line + "\n")
    else:
        # Fallback for local testing
        for line in lines:
            print(line)

    print(f"Wrote: {docx_path}")
    print("Prepared HTML email body with paragraph spacing and superscripts.")

if __name__ == "__main__":
    main()
